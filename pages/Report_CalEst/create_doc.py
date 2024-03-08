from docx import Document
from docx.shared import Inches
from docx.shared import Cm, Pt
from docx.shared import RGBColor
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
import os
import os.path
import json
from datetime import date, datetime, timedelta
import pandas as pd
import glob, os.path
from PIL import Image
import numpy as np




def mk_doc(est,sensor,cod_l,semestre,est_din,est_gapover,est_off,est_ppsd_pic,cod_sen, mes=False):
    
    if mes == True:
        print(semestre[0:4])
        print(semestre[4:len(semestre)])
        anno = semestre[0:4]
        mes_inf=semestre[4:len(semestre)]
        folder_semestre = os.path.dirname(os.path.abspath(__file__))+f"/Inf_mensuales/{est}_{sensor}/{anno}/{mes_inf}/"
    
    if mes == False:
        folder_semestre = os.path.dirname(os.path.abspath(__file__))+f"/Inf_semestrales/{est}_{sensor}/{semestre}/"

    document = Document()
    #fuente
    document.styles["Normal"].font.name= "Arial"
    document.styles["Normal"].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    # Agregar fecha de creación
    today = date.today()
    
    ##ENCABEZADO
    section = document.sections[0]
    header = section.header
    #paragraph = header.paragraphs[0]
    #run = paragraph.add_run()
    #run.add_picture(folder_semestre+'Simbolo_SGC_Color.png', width=Inches(1.25))
    table = header.add_table(rows=1, cols=3, width=Cm(18))
    table.style="Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(0.16)
    table.columns[1].width = Cm(0.35)
    table.columns[2].width = Cm(0.17)
    
    row = table.rows[0]
    cell = row.cells[0]
    cell.add_paragraph().add_run().add_picture(os.path.dirname(os.path.abspath(__file__))+'/Simbolo_SGC_Color.png', width=Inches(1.25))
    cell = row.cells[1]
    
    if cod_l == "30":
        sensor="HH2"
    cod_loc = {"HH":"banda ancha","HN":"aceleración","EH":"corto periodo","BH":"banda ancha","HL":"movimiento fuerte","HH2":"banda intermedia"} #banda ancha HH.00,01..,  corto periodo EH.20,21.., aceleracion HN.10,11..

    if mes == False:
        cell.add_paragraph(f"Informe de funcionamiento y calidad de la Estación de {cod_loc[sensor]} {est} | Semestre {semestre}")
    if mes == True:
        cell.add_paragraph(f"Informe de funcionamiento y calidad de la Estación de {cod_loc[sensor]} {est} | Mes {mes_inf} de {anno}")

    if cod_l == "30": sensor="HH"

    cell.bold = True
    
    #num_pages = len(document.paragraphs)
    
    
    #run.add_text(f"Informe de funcionamiento y calidad de la Estación de {cod_loc[sensor]} {est}")
   #run.element.rPr.rAlign = "right"
   
    ##INFORMACIÓN GENERAL DE LA ESTACIÓN
    file_xlsx = os.path.dirname(os.path.abspath(__file__))+"/dat_est/stationReportCodLocation.xlsx"
        
    tab_est = pd.read_excel(file_xlsx)
    
    #filtro por Estación
    filtro_estacion = (tab_est["IDENTIFICADOR"] == est)
    fil_est = tab_est[filtro_estacion]

    resp_tem, nombre, departamento, municipio, latitud, longitud, altura, resp_elect, resp_tem   = "","","","","","","","",""
    estado, cond_ins, transmi, adquisicion, tip_est, equipos, f_in, f_fi  = "","","","","","","",""
    
    if len(fil_est) > 0:
        
        
        nombre = fil_est["NOMBRE"].iloc[0]
        departamento = fil_est["DEPARTAMENTO"].iloc[0]
        municipio = fil_est["MUNICIPIO"].iloc[0]
        latitud = fil_est["LATITUD (°)"].iloc[0]
        longitud = fil_est["LONGITUD (°)"].iloc[0]
        altura = fil_est["ELEVACION (msnm)"].iloc[0]
        
        resp_elect =fil_est["ELECTRÓNICO RESPONSABLE"].iloc[0]
        resp_tem = fil_est["TEMÁTICO RESPONSABLE"].iloc[0]
        
        estado = fil_est["ESTADO ESTACIÓN"].iloc[0]
        cond_ins = fil_est["CONDICIÓN DE INSTALACIÓN"].iloc[0]
        transmi = fil_est["TIPO DE TRANSMISIÓN"].iloc[0]
        adquisicion = fil_est["TIPO DE ADQUISICIÓN"].iloc[0]
        tip_est = fil_est["TIPO DE ESTACIÓN"].iloc[0]
        
        n_sensores = len(fil_est)
        for e in range(n_sensores):
            
            if np.isnan(fil_est["CÓDIGO LOCALIZACION"].iloc[e]) == False: 
                codigo_loc = int(fil_est["CÓDIGO LOCALIZACION"].iloc[e])
            else:
                codigo_loc = "--"
            
            ig=False
            if str(codigo_loc).rjust(2,"0") == cod_l:
                
                equipos = fil_est["INSTRUMENTACIÓN"].iloc[e]
                f_in = fil_est["FECHA INICIO COD."].iloc[e]
                f_fi = fil_est["FECHA FIN COD."].iloc[e]
                ig = True

        if ig==False:
            equipos = fil_est["INSTRUMENTACIÓN"].iloc[0]
            f_in = fil_est["FECHA INICIO COD."].iloc[0]
            f_fi = fil_est["FECHA FIN COD."].iloc[0]


    
    #print("##Cod localizacion",str(codigo_loc).rjust(2,"0"), cod_l)
    
    ##EXTRAER INFORMACIÓN DE LOS JSON
    list_est_sen = glob.glob(os.path.dirname(os.path.abspath(__file__))+f"/Datos/CM/Tablas_Cal_tiempos/{est}_{sensor}Z*")
    cod_sensor=[]
    if len(list_est_sen) == 2:
        for s in list_est_sen:
            if s[-10:-8] not in cod_sensor:
                cod_sensor.append(s[-10:-8])
    if mes == True:
        folder_semestre = os.path.dirname(os.path.abspath(__file__))+f"/Inf_mensuales/{est}_{sensor}/{anno}/{mes_inf}/"
        ep = "mes"
    if mes == False:
        folder_semestre = os.path.dirname(os.path.abspath(__file__))+f"/Inf_semestrales/{est}_{sensor}/{semestre}/"
        ep = "semestre"

    if os.path.exists(folder_semestre + f"inf_{est}_{sensor}{cod_sen}.json") == True:
        with open(folder_semestre + f"inf_{est}_{sensor}{cod_sen}.json", 'r') as file:
            inf_json = json.load(file)
        
        text_fun= inf_json[0]["funcionamiento"]
        text_dis= inf_json[0]["disponibilidad"]
        text_gap= inf_json[0]["gapover"] 
        text_cal= inf_json[0]["calidad"]
        text_offset= inf_json[0]["offset"]  
        text_ruido=inf_json[0]["ruido"]   
        text_fecha,text_man,text_com,text_vis= inf_json[0]["ultima_v"]  
        text_recom=inf_json[0]["recomendaciones"]   
        text_fech_creacion=inf_json[0]["fecha_creacion"]  
        n_img_ruido= inf_json[0]["n_img_ruido"]
        n_img_recom=inf_json[0]["n_img_recom"]
        
        n_tot_img= inf_json[0]["n_img_ruido"]+inf_json[0]["n_img_recom"]+12   


    #encabezado izquierdo
    #if text_fech_creacion == " ":
    hoy = datetime.now()
    text_fech_creacion= f"{hoy.year}-{hoy.month}-{hoy.day}"

    year=text_fech_creacion.split("-")[0]
    month=text_fech_creacion.split("-")[1].rjust(2,"0")
    day=text_fech_creacion.split("-")[2].rjust(2,"0")
    cell = row.cells[2]
    cell.add_paragraph(f'Sismología \nPor: {resp_tem.title()}\nFecha: {year}/{month}/{day}')
            
    # Agregar título
    paragraph = document.add_paragraph()
    
    run1 = paragraph.add_run()
    for e in range(3): run1.add_break(WD_BREAK.LINE)#salto de linea
    
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f'Estación {nombre.title()} - {est} {sensor}')
    #run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #run.alignment = 1
    run.font.size = Pt(12)
    run.bold = True
    #document.add_heading(f'Estación {nombre} - {est} {sensor}', 1)
    paragraph = document.add_paragraph()
    
    #Agregar información general de la estación
    for e in range(2): run.add_break(WD_BREAK.LINE)
    run2=paragraph.add_run("Departamento: ").bold = True
    run2=paragraph.add_run(f"{departamento.title()}")
    run2=paragraph.add_run("  |  Municipio: ").bold = True
    run2=paragraph.add_run(f"{municipio.title()}")
    run2.add_break(WD_BREAK.LINE)
    run2=paragraph.add_run("Coordenadas de la estación: ").bold = True
    run2=paragraph.add_run(f" Lat. {latitud}, Lon. {longitud}")
    run2.add_break(WD_BREAK.LINE)
    run2=paragraph.add_run("Tipo de transmisión: ").bold = True
    run2=paragraph.add_run(f"{transmi}")
    run2=paragraph.add_run("  |  Tipo de adquisición: ").bold = True
    run2=paragraph.add_run(f"{adquisicion}")
    run2.add_break(WD_BREAK.LINE)
    run2=paragraph.add_run("Condición de instalación: ").bold = True
    run2=paragraph.add_run(f"{cond_ins   }")
    run2=paragraph.add_run("  |  Tipo de estación: ").bold = True
    run2=paragraph.add_run(f"{tip_est}")
    for e in range(2):run2.add_break(WD_BREAK.LINE)
    run2=paragraph.add_run(f"Sensor de {cod_loc[sensor]} - {cod_l} ").bold = True
    run2=paragraph.add_run()
    run2.add_break(WD_BREAK.LINE)
    run2=paragraph.add_run(f"Sensor y digitalizador: ").bold = True
    run2=paragraph.add_run(f"{equipos}")
    run2.add_break(WD_BREAK.LINE)
    run2=paragraph.add_run(f"Fecha inicio: ").bold = True
    run2=paragraph.add_run(f"{f_in}")
    run2=paragraph.add_run("  |  Fecha fin: ").bold = True
    run2=paragraph.add_run(f"{f_fi}")

    
    ##Estado Actual
    paragraph = document.add_paragraph()
    run1 = paragraph.add_run()
    for e in range(2): run1.add_break(WD_BREAK.LINE)#salto de linea
    
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f'Estado actual')
    run.font.size = Pt(12)
    run.bold = True
    paragraph = document.add_paragraph()

    run2=paragraph.add_run("Observación del Sismólogo")
    run2.bold = True
    for e in range(2): run2.add_break(WD_BREAK.LINE)

    file_esta = os.path.dirname(os.path.abspath(__file__))+f"/Estados/{est}.csv"

    if os.path.exists(file_esta) == True:        
    
        tab_est = pd.read_csv(file_esta)
        ultimo = tab_est.tail(1)
        f =ultimo["fecha"].iloc[0] #fecha de actualización
        act_sis = ultimo["act_sis"].iloc[0]
        act_elec =  ultimo["act_elec"].iloc[0]
        f_prob = ultimo["f_prob"].iloc[0]
        e_disp = ultimo['e_disp'].iloc[0]
        p_sist = ultimo["p_sist"].iloc[0]
        if str(act_sis) == "nan" :
            run=paragraph.add_run(" ")
        if str(act_sis) != "nan" :
            run=paragraph.add_run(act_sis)

        
        for e in range(2): run.add_break(WD_BREAK.LINE)
        run2=paragraph.add_run("Observación del Electrónico")
        run2.bold = True
        for e in range(2): run2.add_break(WD_BREAK.LINE)

        if str(act_elec) == "nan" :
            run=paragraph.add_run(" ")
        if str(act_elec) != "nan" :
            run=paragraph.add_run(act_elec)

        
        for e in range(2): run.add_break(WD_BREAK.LINE)
        run2=paragraph.add_run("Fecha del problema: ")
        run2.bold = True
        run=paragraph.add_run(f_prob)
        for e in range(2): run.add_break(WD_BREAK.LINE)
        run2=paragraph.add_run("Estado de disponibilidad: ")
        run2.bold = True
        run=paragraph.add_run(e_disp)
        for e in range(2): run.add_break(WD_BREAK.LINE)
        run2=paragraph.add_run("Problema de sistema: ")
        run2.bold = True
        run=paragraph.add_run(p_sist)
        for e in range(2): run.add_break(WD_BREAK.LINE)










    ##1. FUNCIONAMIENTO
    run3=paragraph.add_run("1. Funcionamiento")
    run3.font.size = Pt(12)
    run3.bold = True
    for e in range(2): run3.add_break(WD_BREAK.LINE)
    #AGREGAR COMENTARIOS DE FUNCIONAMIENTO
    run3=paragraph.add_run(f"{text_fun}")
    for e in range(2): run3.add_break(WD_BREAK.LINE)
    #1.1 AGREGAR COMENTARIOS DE DISPONIBILIDAD
    run3=paragraph.add_run("1.1 Disponibilidad")
    run3.font.size = Pt(12)
    run3.bold = True
    for e in range(2): run3.add_break(WD_BREAK.LINE)
    run3=paragraph.add_run(f"{text_dis}")

    
    
    paragraph2 = document.add_paragraph()
    
    
    
    ##ESTADISTICAS

    if len(est_din[1])!= 0 and len(est_din[1])!= 0:

        #Comportamiento de disponibilidad
        for e in range(1): run3.add_break(WD_BREAK.LINE)

        est_dinz,est_dinn, est_dine=est_din[0],est_din[1],est_din[2]
        min_dis_z,max_dis_z,prom_dis_z,codigo_l=est_dinz #Contiene [min_dis,max_dis,prom_dis,codigo_l]
        min_dis_n,max_dis_n,prom_dis_n,codigo_l=est_dinn #Contiene [min_dis,max_dis,prom_dis,codigo_l]
        min_dis_e,max_dis_e,prom_dis_e,codigo_l=est_dine #Contiene [min_dis,max_dis,prom_dis,codigo_l]


        run3=paragraph2.add_run(f"Comportamiento de disponibilidad en el {ep} para las tres componentes")
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        run3=paragraph2.add_run(f"Z    |    mínimo: {min_dis_z}%,    máximo: {max_dis_z}%,    promedio: {round(prom_dis_z,2)}%")
        run3.add_break(WD_BREAK.LINE)
        run3=paragraph2.add_run(f"N    |    mínimo: {min_dis_n}%,    máximo: {max_dis_n}%,    promedio: {round(prom_dis_n,2)}%")
        run3.add_break(WD_BREAK.LINE)
        run3=paragraph2.add_run(f"E    |    mínimo: {min_dis_e}%,    máximo: {max_dis_e}%,    promedio: {round(prom_dis_e,2)}%")
        
        document.add_page_break()
        
        #GRAFICAS DE DISPONIBILIDAD
        gr=paragraph2.add_run()
        for e in range(2): gr.add_break(WD_BREAK.LINE)
        #gr.add_page_break()
        #gr.add_break(WD_BREAK.PAGE)
        gr.add_picture(folder_semestre+f"disp_{est}_{sensor}Z{cod_sen}.png",width=Inches(6.5)) #disp_ACH1_HHE.png
        gr.add_picture(folder_semestre+f"disp_{est}_{sensor}N{cod_sen}.png",width=Inches(6.5))
        gr.add_picture(folder_semestre+f"disp_{est}_{sensor}E{cod_sen}.png",width=Inches(6.5))
        run3=paragraph2.add_run(f"  Figura 1.")
        run3.font.size = Pt(9)
        run3.bold = True
        run3=paragraph2.add_run(f" Gráfica de disponibilidad en los datos de la estación {est} en sus tres componentes.")
        run3.font.size = Pt(9)
        

        
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        
        #1.2 AGREGAR COMENTARIOS DE GAPS Y OVERLAPS
        run3=paragraph2.add_run("1.2 Gaps y Overlaps")
        run3.font.size = Pt(12)
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        run3=paragraph2.add_run(f"{text_gap}")
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        paragraph3 = document.add_paragraph()

        ##ESTADISTICAS
        est_gapoverz,est_gapovern, est_gapovere=est_gapover[0],est_gapover[1],est_gapover[2]
        num_gaps_z,max_gaps_z,prom_gaps_z,num_over_z, max_over_z, prom_over_z,codigo_l_over=est_gapoverz #Contiene [num_gaps,max_gaps,gaps_prom,num_overlaps,max_overlap,overlaps_prom, codigo_l]
        num_gaps_n,max_gaps_n,prom_gaps_n,num_over_n, max_over_n, prom_over_n,codigo_l_over=est_gapovern #Contiene [num_gaps,max_gaps,gaps_prom,num_overlaps,max_overlap,overlaps_prom, codigo_l]
        num_gaps_e,max_gaps_e,prom_gaps_e,num_over_e, max_over_e, prom_over_e,codigo_l_over=est_gapovere #Contiene [num_gaps,max_gaps,gaps_prom,num_overlaps,max_overlap,overlaps_prom, codigo_l]

        run3=paragraph3.add_run(f"Comportamiento de gaps y overlaps en el {ep} para las tres componentes")
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run("Gaps")
        run3.bold = True
        for e in range(1): run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"Z    |    número de gaps: {num_gaps_z},    máximo: {max_gaps_z},    promedio: {round(prom_gaps_z,2)}")
        run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"N    |    número de gaps: {num_gaps_n},    máximo: {max_gaps_n},    promedio: {round(prom_gaps_n,2)}")
        run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"E    |    número de gaps: {num_gaps_e},    máximo: {max_gaps_e},    promedio: {round(prom_gaps_e,2)}")
        for e in range(2): run3.add_break(WD_BREAK.LINE)

        run3=paragraph3.add_run("Overlaps")
        run3.bold = True
        for e in range(1): run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"Z    |    número de overlaps: {num_over_z},    máximo: {max_over_z},    promedio: {round(prom_over_z,2)}")
        run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"N    |    número de overlaps: {num_over_n},    máximo: {max_over_n},    promedio: {round(prom_over_n,2)}")
        run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"E    |    número de overlaps: {num_over_e},    máximo: {max_over_e},    promedio: {round(prom_over_e,2)}")
        #for e in range(2): run3.add_break(WD_BREAK.LINE)
        document.add_page_break()
        
        
        #GRAFICAS DE GAPS OVERLAPS
        gr2=paragraph3.add_run()
        for e in range(2): gr2.add_break(WD_BREAK.LINE)
        #gr.add_page_break()
        #gr.add_break(WD_BREAK.PAGE)
        gr2.add_picture(folder_semestre+f"gapover_{est}_{sensor}Z{cod_sen}.png",width=Inches(6.5)) #disp_ACH1_HHE.png
        gr2.add_picture(folder_semestre+f"gapover_{est}_{sensor}N{cod_sen}.png",width=Inches(6.5))
        gr2.add_picture(folder_semestre+f"gapover_{est}_{sensor}E{cod_sen}.png",width=Inches(6.5))
        run4=paragraph3.add_run(f"  Figura 2.")
        run4.bold = True
        run4.font.size = Pt(9)
        run4=paragraph3.add_run(f"  Gráfica de gap y overlaps en los datos de la estación {est} en sus tres componentes.")
        run4.font.size = Pt(9)
        for e in range(2): run4.add_break(WD_BREAK.LINE)
        
        
        
        ##2. CALIDAD
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run("2. Calidad")
        run3.font.size = Pt(12)
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        #AGREGAR COMENTARIOS DE CALIDAD
        run3=paragraph3.add_run(f"{text_cal}")
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        
        #1.1 AGREGAR COMENTARIOS DE DISPONIBILIDAD
        run3=paragraph3.add_run("2.1 Offset")
        run3.font.size = Pt(12)
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"{text_offset}")

        for e in range(2): run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"Comportamiento de offset en el {ep} para las tres componentes")
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        
        ##ESTADISTICAS
        est_offz,est_offn, est_offe=est_off[0],est_off[1],est_off[2]
        min_offsetz, max_offsetz, offset_promz,codigo_l_off = est_offz #contiene [min_offset,max_offset,offset_prom, codigo_l]
        min_offsetn, max_offsetn, offset_promn,codigo_l_off = est_offn
        min_offsete, max_offsete, offset_prome,codigo_l_off = est_offe
        run3=paragraph3.add_run(f"Z    |    mínimo: {min_offsetz},    máximo: {max_offsetz},    promedio: {round(offset_promz,2)}")
        run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"N    |    mínimo: {min_offsetn},    máximo: {max_offsetn},    promedio: {round(offset_promn,2)}")
        run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"E    |    mínimo: {min_offsete},    máximo: {max_offsete},    promedio: {round(offset_prome,2)}")
        for e in range(2): run3.add_break(WD_BREAK.LINE)

        
        #GRAFICAS DE OFFSET
        gr3=paragraph3.add_run()
        for e in range(1): gr3.add_break(WD_BREAK.LINE)
        #gr.add_page_break()
        #gr.add_break(WD_BREAK.PAGE)
        gr3.add_picture(folder_semestre+f"offset_{est}_{sensor}Z{cod_sen}.png",width=Inches(6.5)) #disp_ACH1_HHE.png
        gr3.add_picture(folder_semestre+f"offset_{est}_{sensor}N{cod_sen}.png",width=Inches(6.5))
        gr3.add_picture(folder_semestre+f"offset_{est}_{sensor}E{cod_sen}.png",width=Inches(6.5))
        run4=paragraph3.add_run(f"  Figura 3.")
        run4.font.size = Pt(9)
        run4.bold = True
        run4=paragraph3.add_run(f"  Gráfica de offset en los datos de la estación {est} en sus tres componentes.")
        run4.font.size = Pt(9)
        for e in range(2): run4.add_break(WD_BREAK.LINE)
        
        
        
        #2.2 AGREGAR COMENTARIOS DE DISPONIBILIDAD
        run3=paragraph3.add_run("2.2 Análisis de ruido")
        run3.font.size = Pt(12)
        run3.bold = True
        
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        #%PPSC
        run3=paragraph3.add_run("Porcentaje fuera de las curvas de Peterson de la media del espectro probabilístico de densidad de potencia (%PPSD) y picos")
        run3.font.size = Pt(12)
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        
        if cod_l == "30":
            sensor="HH2"
            
        t_ppsd_sens={"HH":"Para las estaciones de banda ancha este %ppsd se espera que esté alrededor del 0% y esto nos dirá que las frecuencias registradas se encuentran dentro de lo normal o no.",
        "HH2":"Para las estaciones de banda intermedia este %ppsd se espera que esté alrededor del 0% y esto nos dirá que las frecuencias registradas se encuentran dentro de lo normal o no.",
        "BH":"Para las estaciones de banda intermedia este %ppsd se espera que esté alrededor del 0% y esto nos dirá que las frecuencias registradas se encuentran dentro de lo normal o no.",
        "HN":"Para las estaciones de aceleración este %ppsd se espera que esté alrededor del 30% y esto nos dirá que las frecuencias registradas se encuentran dentro de lo normal o no.",
        "EH":"Para las estaciones de corto periodo este %ppsd se espera que esté alrededor del 25% y esto nos dirá que las frecuencias registradas se encuentran dentro de lo normal o no. ",
        "HL":"Para las estaciones de movimiento fuerte este %ppsd se espera que esté alrededor del 30% - 40% y esto nos dirá que las frecuencias registradas se encuentran dentro de lo normal o no.",
        }
        if cod_l == "30": sensor="HH"
        run3=paragraph3.add_run(f"El %ppsd es el porcentaje de cuánto de la media del espectro de ruido de la estación se encuentra por fuera de las curvas de Peterson, {t_ppsd_sens[sensor]}")
        for e in range(2): run3.add_break(WD_BREAK.LINE)

        ##ESTADISTICAS PPSD Y PICOS
        run3=paragraph3.add_run(f"Comportamiento del %ppsd y picos en el {ep} para las tres componentes.")
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        
        est_ppsd_picz,est_ppsd_picn, est_ppsd_pice=est_ppsd_pic[0],est_ppsd_pic[1],est_ppsd_pic[2]
        ppsd_prom_z,num_picos_z,max_picos_z,codigo_l_pp=est_ppsd_picz #Contiene [ppsd_prom,num_picos,max_picos, codigo_l]
        ppsd_prom_n,num_picos_n,max_picos_n,codigo_l_pp=est_ppsd_picn #Contiene [ppsd_prom,num_picos,max_picos, codigo_l]
        ppsd_prom_e,num_picos_e,max_picos_e,codigo_l_pp=est_ppsd_pice #Contiene [ppsd_prom,num_picos,max_picos, codigo_l]

        
        run3=paragraph3.add_run(f"Z    |    promedio %ppsd: {round(ppsd_prom_z,2)},    número de picos: {num_picos_z},    máximo de picos: {max_picos_z}")
        run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"N    |    promedio %ppsd: {round(ppsd_prom_n,2)},    número de picos: {num_picos_n},    máximo de picos: {max_picos_n}")
        run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"E    |    promedio %ppsd: {round(ppsd_prom_e,2)},    número de picos: {num_picos_e},    máximo de picos: {max_picos_e}")
        for e in range(2): run3.add_break(WD_BREAK.LINE)

        #GRAFICAS DE %ppsd y picos
        gr4=paragraph3.add_run()
        for e in range(1): gr3.add_break(WD_BREAK.LINE)
        #gr.add_page_break()
        #gr.add_break(WD_BREAK.PAGE)
        gr4.add_picture(folder_semestre+f"ppsd_{est}_{sensor}Z{cod_sen}.png",width=Inches(6.5)) # ppsd_ACH1_HHZ.png
        gr4.add_picture(folder_semestre+f"ppsd_{est}_{sensor}N{cod_sen}.png",width=Inches(6.5))
        gr4.add_picture(folder_semestre+f"ppsd_{est}_{sensor}E{cod_sen}.png",width=Inches(6.5))
        run4=paragraph3.add_run(f"  Figura 4.")
        run4.font.size = Pt(9)
        run4.bold = True
        run4=paragraph3.add_run(f" Gráfica de %ppsd y picos en los datos de la estación {est} en sus tres componentes.")
        run4.font.size = Pt(9)
        for e in range(2): run4.add_break(WD_BREAK.LINE)
        
        

    if len(est_din[1]) == 0 and len(est_din[1]) == 0:
        
        #ESTADISTICAS
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        est_dinz,est_dinn, est_dine=est_din[0],est_din[1],est_din[2]
        min_dis_z,max_dis_z,prom_dis_z,codigo_l=est_dinz #Contiene [min_dis,max_dis,prom_dis,codigo_l]
        run3=paragraph2.add_run(f"Comportamiento de disponibilidad en el {ep} para la componente Z")
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        run3=paragraph2.add_run(f"Z    |    mínimo: {min_dis_z}%,    máximo: {max_dis_z}%,    promedio: {round(prom_dis_z,2)}%")

        #GRAFICAS DE DISPONIBILIDAD
        gr=paragraph2.add_run()
        for e in range(2): gr.add_break(WD_BREAK.LINE)
        #gr.add_page_break()
        #gr.add_break(WD_BREAK.PAGE)
        gr.add_picture(folder_semestre+f"disp_{est}_{sensor}Z{cod_sen}.png",width=Inches(6.5)) #disp_ACH1_HHE.png
        run3=paragraph2.add_run(f"  Figura 1.")
        run3.font.size = Pt(9)
        run3.bold = True
        run3=paragraph2.add_run(f" Gráfica de disponibilidad en los datos de la estación {est} en la componente Z.")
        run3.font.size = Pt(9)
        
        


        
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        
        #1.2 AGREGAR COMENTARIOS DE GAPS Y OVERLAPS
        run3=paragraph2.add_run("1.2 Gaps y Overlaps")
        run3.font.size = Pt(12)
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        run3=paragraph2.add_run(f"{text_gap}")
        
        paragraph3 = document.add_paragraph()
        for e in range(2): run3.add_break(WD_BREAK.LINE)

        ##ESTADISTICAS
        est_gapoverz,est_gapovern, est_gapovere=est_gapover[0],est_gapover[1],est_gapover[2]
        num_gaps_z,max_gaps_z,prom_gaps_z,num_over_z, max_over_z, prom_over_z,codigo_l_over=est_gapoverz #Contiene [num_gaps,max_gaps,gaps_prom,num_overlaps,max_overlap,overlaps_prom, codigo_l]

        run3=paragraph3.add_run(f"Comportamiento de gaps y overlaps en el {ep} para la componente Z")
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run("Gaps")
        run3.bold = True
        for e in range(1): run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"Z    |    número de gaps: {num_gaps_z},    máximo: {max_gaps_z},    promedio: {round(prom_gaps_z,2)}")
        for e in range(2): run3.add_break(WD_BREAK.LINE)

        run3=paragraph3.add_run("Overlaps")
        run3.bold = True
        for e in range(1): run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"Z    |    número de overlaps: {num_over_z},    máximo: {max_over_z},    promedio: {round(prom_over_z,2)}")
        
        #document.add_page_break()
        
        paragraph3 = document.add_paragraph()
        
        #GRAFICAS DE GAPS OVERLAPS
        gr2=paragraph3.add_run()
        for e in range(1): gr2.add_break(WD_BREAK.LINE)
        #gr.add_page_break()
        #gr.add_break(WD_BREAK.PAGE)
        gr2.add_picture(folder_semestre+f"gapover_{est}_{sensor}Z{cod_sen}.png",width=Inches(6.5)) #disp_ACH1_HHE.png
        run4=paragraph3.add_run(f"  Figura 2. ")
        run4.font.size = Pt(9)
        run4.bold = True
        run4=paragraph3.add_run(f" Gráfica de gap y overlaps en los datos de la estación {est} en la componente Z.")
        run4.font.size = Pt(9)
        for e in range(2): run4.add_break(WD_BREAK.LINE)
        
        
        #for e in range(2): run3.add_break(WD_BREAK.LINE)
        
        ##2. CALIDAD
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run("2. Calidad")
        run3.font.size = Pt(12)
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        #AGREGAR COMENTARIOS DE CALIDAD
        run3=paragraph3.add_run(f"{text_cal}")

        for e in range(1): run3.add_break(WD_BREAK.LINE)
        
        
        #1.1 AGREGAR COMENTARIOS DE DISPONIBILIDAD
        run3=paragraph3.add_run("2.1 Offset")
        run3.font.size = Pt(12)
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"{text_offset}")
        
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        run3=paragraph3.add_run(f"Comportamiento de offset en el {ep} para las tres componentes")
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        ##ESTADISTICAS
        est_offz,est_offn, est_offe=est_off[0],est_off[1],est_off[2]
        min_offsetz, max_offsetz, offset_promz,codigo_l_off = est_offz #contiene [min_offset,max_offset,offset_prom, codigo_l]

        run3=paragraph3.add_run(f"Z    |    mínimo: {min_offsetz},    máximo: {max_offsetz},    promedio: {round(offset_promz,2)}")

        for e in range(1): run3.add_break(WD_BREAK.LINE)
        
        #GRAFICAS DE OFFSET
        gr3=paragraph3.add_run()
        for e in range(1): gr3.add_break(WD_BREAK.LINE)
        #gr.add_page_break()
        #gr.add_break(WD_BREAK.PAGE)
        gr3.add_picture(folder_semestre+f"offset_{est}_{sensor}Z{cod_sen}.png",width=Inches(6.5)) #disp_ACH1_HHE.png

        run4=paragraph3.add_run(f"  Figura 3.")
        run4.font.size = Pt(9)
        run4.bold = True
        run4=paragraph3.add_run(f" Gráfica de offset en los datos de la estación {est} en la componente Z.")
        run4.font.size = Pt(9)
        for e in range(2): run4.add_break(WD_BREAK.LINE)
        
        run3=paragraph3.add_run(f"Comportamiento de offset en el {ep} para la componente Z")
        run3.bold = True
        
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        
        #2.2 AGREGAR COMENTARIOS DE DISPONIBILIDAD
        run3=paragraph3.add_run("2.2 Análisis de ruido")
        run3.font.size = Pt(12)
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        #%PPSC
        run3=paragraph3.add_run("Porcentaje fuera de las curvas de Peterson de la media del espectro probabilístico de densidad de potencia (%PPSD) y picos")
        run3.font.size = Pt(12)
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        
        if cod_l == "30":
            sensor="HH2"
            
        t_ppsd_sens={"HH":"Para las estaciones de banda ancha este %ppsd se espera que esté alrededor del 0% y esto nos dirá que las frecuencias registradas se encuentran dentro de lo normal o no.",
        "HH2":"Para las estaciones de banda intermedia este %ppsd se espera que esté alrededor del 0% y esto nos dirá que las frecuencias registradas se encuentran dentro de lo normal o no.",
        "BH":"Para las estaciones de banda intermedia este %ppsd se espera que esté alrededor del 0% y esto nos dirá que las frecuencias registradas se encuentran dentro de lo normal o no.",
        "HN":"Para las estaciones de aceleración este %ppsd se espera que esté alrededor del 30% y esto nos dirá que las frecuencias registradas se encuentran dentro de lo normal o no.",
        "EH":"Para las estaciones de corto periodo este %ppsd se espera que esté alrededor del 25% y esto nos dirá que las frecuencias registradas se encuentran dentro de lo normal o no. ",
        "HL":"Para las estaciones de movimiento fuerte este %ppsd se espera que esté alrededor del 30% - 40% y esto nos dirá que las frecuencias registradas se encuentran dentro de lo normal o no.",
        }
        run3=paragraph3.add_run(f"El %ppsd es el porcentaje de cuánto de la media del espectro de ruido de la estación se encuentra por fuera de las curvas de Peterson, {t_ppsd_sens[sensor]}")
        for e in range(2): run3.add_break(WD_BREAK.LINE)

        run3=paragraph3.add_run(f"Comportamiento del %ppsd y picos en el {ep} para la componente Z.")
        run3.bold = True
        for e in range(2): run3.add_break(WD_BREAK.LINE)
        
        ##ESTADISTICAS PPSD Y PICOS
        est_ppsd_picz,est_ppsd_picn, est_ppsd_pice=est_ppsd_pic[0],est_ppsd_pic[1],est_ppsd_pic[2]
        ppsd_prom_z,num_picos_z,max_picos_z,codigo_l_pp=est_ppsd_picz #Contiene [ppsd_prom,num_picos,max_picos, codigo_l]

        
        run3=paragraph3.add_run(f"Z    |    promedio %ppsd: {round(ppsd_prom_z,2)},    número de picos: {num_picos_z},    máximo de picos: {max_picos_z}")
        for e in range(2): run3.add_break(WD_BREAK.LINE)

        #GRAFICAS DE %ppsd y picos
        gr4=paragraph3.add_run()
        for e in range(1): gr3.add_break(WD_BREAK.LINE)
        #gr.add_page_break()
        #gr.add_break(WD_BREAK.PAGE)
        gr4.add_picture(folder_semestre+f"ppsd_{est}_{sensor}Z{cod_sen}.png",width=Inches(6.5)) # ppsd_ACH1_HHZ.png

        run4=paragraph3.add_run(f"  Figura 4.")
        run4.font.size = Pt(9)
        run4.bold = True
        run4=paragraph3.add_run(f" Gráfica de %ppsd y picos en los datos de la estación {est} en la componente Z.")
        run4.font.size = Pt(9)
        for e in range(2): run4.add_break(WD_BREAK.LINE)
        
        
    
    
    #Espectro 
    run3=paragraph3.add_run("Espectro")
    run3.font.size = Pt(12)
    run3.bold = True
    for e in range(2): run3.add_break(WD_BREAK.LINE)
    
    run3=paragraph3.add_run(f"{text_ruido}")
    for e in range(2): run3.add_break(WD_BREAK.LINE)
    
    #GRAFICAS DE ESPECTROS 
    for e in range(n_img_ruido):
        #gr4=paragraph3.add_run()
        print("#",folder_semestre, cod_sen,e+1)
        print("#",glob.glob(folder_semestre+f"img_ruido{cod_sen}_{e+1}*"))
        img_ruido = glob.glob(folder_semestre+f"img_ruido{cod_sen}_{e+1}*")[0].split("/")[-1]
        if img_ruido != f"img_ruido{cod_sen}_{e+1}.png":
            image = Image.open(folder_semestre+img_ruido)
            image.save(folder_semestre+f"img_ruido{cod_sen}_{e+1}.png")
        run4=paragraph3.add_run().add_picture(folder_semestre+f"img_ruido{cod_sen}_{e+1}.png",width=Inches(5))
        run4=paragraph3.add_run(f"\nFigura {e+1+4}.")
        run4.font.size = Pt(9)
        run4.bold = True
        run4=paragraph3.add_run(f" Espectro de ruido en los datos de la estación {est}.\n\n")
        run4.font.size = Pt(9)

    
    
    ##3. ULTIMA VISITA
    for e in range(2): run3.add_break(WD_BREAK.LINE)
    run3=paragraph3.add_run("\n3. Última visita")
    run3.font.size = Pt(12)
    run3.bold = True
    for e in range(2): run3.add_break(WD_BREAK.LINE)
    #AGREGAR COMENTARIOS DE ULTIMA VISITA
    texto_completo=f"La última visita a la estación fue el {text_fecha} por {text_vis.title()} realizando  {text_man.lower()}, {text_com.lower()}"
    run3=paragraph3.add_run(f"{texto_completo}")
    for e in range(2): run3.add_break(WD_BREAK.LINE)
    
    ##4. RECOMENDACIONES
    for e in range(2): run3.add_break(WD_BREAK.LINE)
    run3=paragraph3.add_run("4. Recomendaciones")
    run3.font.size = Pt(12)
    run3.bold = True
    for e in range(2): run3.add_break(WD_BREAK.LINE)
    run3=paragraph3.add_run(f"{text_recom}")
    for e in range(2): run3.add_break(WD_BREAK.LINE)
    
    #GRAFICAS DE RECOMENDACIONES
    for reco in range(n_img_recom):
        img_ruido = glob.glob(folder_semestre+f"img_recom{cod_sen}_{reco+1}*")[0].split("/")[-1]
        if img_ruido != f"img_recom{cod_sen}_{reco+1}.png":
            image = Image.open(folder_semestre+img_ruido)
            image.save(folder_semestre+f"img_recom{cod_sen}_{reco+1}.png")
        run4=paragraph3.add_run().add_picture(folder_semestre+f"img_recom{cod_sen}_{reco+1}.png",width=Inches(5))
        run4=paragraph3.add_run(f"\nFigura {n_img_ruido+reco+1+4}. ")
        run4.font.size = Pt(9)
        run4.bold = True
        run4=paragraph3.add_run(f" Imagen de apoyo de la estación {est}.\n\n")
        run4.font.size = Pt(9)
 
    

    document.save(folder_semestre+f'{est}_{cod_l}_{semestre}_informe.docx')

def mk_doc_all(est,sensor,cod_l,semestre,est_din,est_gapover,est_off,est_ppsd_pic,cod_sen, mes=False):

    if mes == True:
        print(semestre[0:4])
        print(semestre[4:len(semestre)])
        anno = semestre[0:4]
        mes_inf=semestre[4:len(semestre)]
        folder_semestre = os.path.dirname(os.path.abspath(__file__))+f"/Inf_mensuales/{est}_{sensor}/{anno}/{mes_inf}/"
    
    if mes == False:
        folder_semestre = os.path.dirname(os.path.abspath(__file__))+f"/Inf_semestrales/{est}_{sensor}/{semestre}/"

    document = Document()
    print("")
#mk_doc("ACH1","HH","00","2023-I")